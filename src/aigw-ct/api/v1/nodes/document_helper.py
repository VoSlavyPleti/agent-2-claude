import re
import json
from copy import deepcopy
from io import BytesIO
from typing import List
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from aigw_ct.context import APP_CTX

logger = APP_CTX.get_logger()

class DocumentProcessingHelpers:
    """Класс для помощи в обработке документов"""

    """
    Основные методы:
        extract_text: Извлечение текста документа из байтового массива
        precise_trim_docx: Разделение документа на текст и формы по регулярному выражению
        precise_trim_forms: Извлечение форм из документа по регулярным выражениям начала и конца в новый байт-массив
    """

    def __init__(self, bytes_array, split_phrase=None):
        self.bytes_array = bytes_array
        self.split_phrase = split_phrase
        bitt = bytes(int(b) for b in self.bytes_array)
        self.doc = Document(BytesIO(bitt))

    def _contains_text_forms(self, element, text):
        """Проверка наличия текста в элементе (fuzzy-matching через _clear_text)"""
        if isinstance(element, Paragraph):
            return self._clear_text(text) in self._clear_text(element.text)
        elif isinstance(element, Table):
            for row in element.rows:
                for cell in row.cells:
                    if self._clear_text(text) in self._clear_text(cell.text):
                        return True
        return False

    def _find_element_by_marker(self, all_elements, marker):
        """
        3-уровневый поиск маркера среди элементов документа.
        Level 1: exact match через _clear_text (текущее поведение)
        Level 2: normalized — нормализация пробелов и регистра
        Level 3: word overlap — ≥70% слов маркера присутствуют в элементе
        Возвращает элемент или None.
        """
        # Level 1: exact (через _clear_text — уже убирает спецсимволы)
        for element, obj in all_elements:
            if self._contains_text_forms(obj, marker):
                return element

        # Level 2: normalized — более мягкая очистка
        norm_marker = re.sub(r'\s+', ' ', marker.lower().strip())
        for element, obj in all_elements:
            if isinstance(obj, Paragraph):
                norm_text = re.sub(r'\s+', ' ', obj.text.lower().strip())
                if norm_marker in norm_text:
                    return element
            elif isinstance(obj, Table):
                for row in obj.rows:
                    for cell in row.cells:
                        norm_text = re.sub(r'\s+', ' ', cell.text.lower().strip())
                        if norm_marker in norm_text:
                            return element

        # Level 3: word overlap ≥70%
        words = [w for w in norm_marker.split() if len(w) > 2]
        if not words:
            return None
        threshold = 0.7
        for element, obj in all_elements:
            text = ""
            if isinstance(obj, Paragraph):
                text = obj.text.lower()
            elif isinstance(obj, Table):
                text = " ".join(cell.text.lower() for row in obj.rows for cell in row.cells)
            matched = sum(1 for w in words if w in text)
            if matched / len(words) >= threshold:
                logger.info(f"[precise_trim_forms] Fuzzy match: {matched}/{len(words)} words for marker '{marker[:50]}'")
                return element

        return None

    def _clear_text(self, text):
        """Очистка текста от всего кроме букв и цифр"""
        text = text.lower()
        text = text.strip()
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '', text)
        return text

    def extract_text(self) -> str:
        doc = self.doc
        result = []
        for el in doc.iter_inner_content():

            if el._element.tag.endswith('p'):
                if el.text != '':
                    text = el.text.strip()
                    text = re.sub(r'\s{3,}', '\n', text)
                    result.append(text)
                    result.append('\n')
            elif el._element.tag.endswith('tbl'):
                table = []
                result.append('******Начало таблицы******\n')
                for row in el.rows:
                    for cell in row.cells:
                        text = cell.text
                        text = re.sub('\n', ' ', text)
                        text = re.sub(r'\s+', ' ', text)
                        if text == '':
                            table.append('Пустая ячейка')
                            continue
                        table.append(text.strip())
                    table.append('\n')
                result.append(f"|{'|'.join(table)}")
                result.append('******Конец таблицы******\n\n')
        return ''.join(result)

    def precise_trim_docx(self) -> tuple([str, List[bytes]]):
        """
        Удаляет ВСЕ элементы документа до элемента с start_text
        и ВСЕ элементы после элемента с end_text.
        Сохраняет результат в новый файл.
        """
        doc = self.doc
        start_text = self.split_phrase.lower()
        start_text = re.sub("o", "о", start_text)
        start_text = re.sub(r'[^\S\n]+', ' ', start_text)
        body = doc._element.body

        # Находим начальный и конечный элементы
        start_element = None
        end_element = body[-2]

        # Запись разметки исходного документа
        original_section = doc.sections[0]
        original_left_margin = original_section.left_margin  # Отступ слева
        original_right_margin = original_section.right_margin  # Отступ справа
        original_top_margin = original_section.top_margin  # Отступ сверху
        original_bottom_margin = original_section.bottom_margin  # Отступ снизу
        original_orientation = original_section.orientation  # Ориентация страницы

        # Ищем все элементы документа
        all_elements = []
        for element in body.iterchildren():
            if element.tag.endswith('p'):  # Параграф
                all_elements.append((element, Paragraph(element, doc)))
            elif element.tag.endswith('tbl'):  # Таблица
                all_elements.append((element, Table(element, doc)))

        # Ищем начальный маркер
        for element, obj in all_elements:
            if self._contains_text_forms(obj, start_text):
                start_element = element
                # break  # для первого вхождения

        if start_element is None:
            buffer = BytesIO()
            doc.save(buffer)
            byte_array = buffer.getvalue()
            buffer.close()
            return ("failed", list(byte_array))


        # Находим индексы наших элементов
        all_elements_only = [e[0] for e in all_elements]
        start_idx = all_elements_only.index(start_element)
        end_idx = all_elements_only.index(end_element)

        # Удаляем элементы до начального маркера
        for element in all_elements_only[:start_idx]:
            body.remove(element)

        doc.add_section()
        new_section = doc.sections[0]
        new_section.left_margin = original_left_margin  # Отступ слева
        new_section.right_margin = original_right_margin  # Отступ справа
        new_section.top_margin = original_top_margin  # Отступ сверху
        new_section.bottom_margin = original_bottom_margin  # Отступ снизу
        new_section.orientation = original_orientation  # Ориентация страницы
        buffer = BytesIO()
        doc.save(buffer)
        byte_array = buffer.getvalue()
        buffer.close()
        return ("processed", list(byte_array))

    def precise_trim_forms(self, start_text: str, end_text: str):
        """
        Удаляет ВСЕ элементы документа до элемента с start_text
        и ВСЕ элементы после элемента с end_text.
        Сохраняет результат в новый файл.

        Использует 3-уровневый поиск маркеров (exact → normalized → word overlap).

        :param start_text: текст начального маркера
        :param end_text: текст конечного маркера
        """
        doc = self.doc
        body = doc._element.body

        # Запись разметки исходного документа
        original_section = doc.sections[0]
        original_left_margin = original_section.left_margin
        original_right_margin = original_section.right_margin
        original_top_margin = original_section.top_margin
        original_bottom_margin = original_section.bottom_margin
        original_orientation = original_section.orientation

        # Ищем все элементы документа
        all_elements = []
        for element in body.iterchildren():
            if element.tag.endswith('p'):
                all_elements.append((element, Paragraph(element, doc)))
            elif element.tag.endswith('tbl'):
                all_elements.append((element, Table(element, doc)))

        # Ищем начальный маркер (3-уровневый поиск)
        start_element = self._find_element_by_marker(all_elements, start_text)

        if start_element is None:
            raise ValueError(f"Не удалось найти начальный маркер '{start_text}' в документе")

        # Ищем конечный маркер (только после start_element, 3-уровневый поиск)
        start_idx = [e[0] for e in all_elements].index(start_element)
        elements_after_start = all_elements[start_idx + 1:]
        end_element = self._find_element_by_marker(elements_after_start, end_text)

        if end_element is None:
            raise ValueError(f"Не удалось найти конечный маркер '{end_text}' в документе")

        # Находим индексы наших элементов
        all_elements_only = [e[0] for e in all_elements]
        start_idx = all_elements_only.index(start_element)
        end_idx = all_elements_only.index(end_element)

        # Удаляем элементы до начального маркера
        for element in all_elements_only[:start_idx]:
            body.remove(element)

        # Удаляем элементы после конечного маркера
        remaining_elements = list(body.iterchildren())
        end_idx_in_remaining = remaining_elements.index(end_element)

        for element in remaining_elements[end_idx_in_remaining+1:]:
            body.remove(element)

        # Применение исходной разметки документа
        doc.add_section()
        new_section = doc.sections[0]
        new_section.left_margin = original_left_margin  # Отступ слева
        new_section.right_margin = original_right_margin  # Отступ справа
        new_section.top_margin = original_top_margin  # Отступ сверху
        new_section.bottom_margin = original_bottom_margin  # Отступ снизу
        new_section.orientation = original_orientation  # Ориентация страницы
        buffer = BytesIO()
        doc.save(buffer)
        byte_array = buffer.getvalue()
        buffer.close()
        return list(byte_array)

class FormFiller:
    """
    Заполнение форм данными.

    Улучшения:
    - XML-level заполнение через lxml для максимального сохранения форматирования
    - Расширенный паттерн плейсхолдеров (подчёркивания, точки, скобки)
    - Работа на уровне run'ов для сохранения форматирования (шрифт, размер, жирность)
    - Улучшенный поиск пропусков в соседних run'ах
    - Копирование стилей при заполнении таблиц
    - Лучшая обработка ошибок с логированием
    """

    # Расширенный паттерн для различных типов плейсхолдеров
    PLACEHOLDER_PATTERN = r'__{2,}|\[_+\]|<_+>|\.{5,}'

    @staticmethod
    def _fill_run_xml(paragraph, pattern, value):
        """
        XML-level заполнение: находит run с плейсхолдером, заменяет текст,
        полностью сохраняет rPr (форматирование).
        Возвращает True если замена произошла.
        """
        for run_elem in paragraph._element.findall(qn('w:r')):
            t_elem = run_elem.find(qn('w:t'))
            if t_elem is not None and t_elem.text and re.search(pattern, t_elem.text):
                t_elem.text = re.sub(pattern, value, t_elem.text, count=1)
                t_elem.set(qn('xml:space'), 'preserve')
                return True
        return False

    @staticmethod
    def _detect_field_codes(paragraph):
        """Обнаруживает field codes (<w:fldChar> и <w:fldSimple>) в параграфе."""
        fld_chars = paragraph._element.findall('.//' + qn('w:fldChar'))
        fld_simples = paragraph._element.findall('.//' + qn('w:fldSimple'))
        return len(fld_chars) > 0 or len(fld_simples) > 0

    @staticmethod
    def _fill_field_code(paragraph, value):
        """
        Заполняет MERGEFIELD/FORMTEXT: находит result run между fldChar separate и end.
        Возвращает True если заполнение произошло.
        """
        runs = list(paragraph._element.findall(qn('w:r')))
        in_result = False
        for run in runs:
            fld_char = run.find(qn('w:fldChar'))
            if fld_char is not None:
                fld_type = fld_char.get(qn('w:fldCharType'))
                if fld_type == 'separate':
                    in_result = True
                elif fld_type == 'end':
                    in_result = False
            elif in_result:
                t = run.find(qn('w:t'))
                if t is not None:
                    t.text = value
                    t.set(qn('xml:space'), 'preserve')
                    return True
        return False

    @staticmethod
    def _get_real_cell_index(row, cell):
        """Возвращает реальный индекс ячейки с учётом объединённых ячеек (gridSpan)."""
        idx = 0
        for c in row.cells:
            if c._element == cell._element:
                return idx
            tc_pr = c._element.find(qn('w:tcPr'))
            if tc_pr is not None:
                grid_span = tc_pr.find(qn('w:gridSpan'))
                if grid_span is not None:
                    idx += int(grid_span.get(qn('w:val'), 1))
                else:
                    idx += 1
            else:
                idx += 1
        return idx

    @staticmethod
    def _fill_table_cell_xml(cell, value, source_cell=None):
        """
        XML-level заполнение ячейки таблицы: добавляет run с текстом,
        копируя rPr из source_cell если доступен.
        """
        p = cell.paragraphs[0]
        if p.runs:
            # Есть run — заполняем первый
            p.runs[0].text = str(value)
            return True

        # Нет runs — создаём через XML с копированием стиля
        new_run = p.add_run(str(value))
        if source_cell and source_cell.paragraphs and source_cell.paragraphs[-1].runs:
            source_rpr = source_cell.paragraphs[-1].runs[0]._element.find(qn('w:rPr'))
            if source_rpr is not None:
                new_run._element.insert(0, deepcopy(source_rpr))
        return True

    @staticmethod
    def _clear_text(text):
        """Очистка текста от всего кроме букв и цифр"""
        text = text.lower()
        text = text.strip()
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^a-zA-Zа-яА-ЯёЁ0-9]', '', text)
        return text

    @staticmethod
    def copy_styles(current_cell, previous_cell):
        """Копирование стилей из предыдущего параграфа/ячейки в текущую"""
        if current_cell and previous_cell:
            try:
                current_cell.style = previous_cell.style
                if previous_cell.runs and current_cell.runs:
                    source_font = previous_cell.runs[0].font
                    target_font = current_cell.runs[-1].font

                    target_font.name = source_font.name
                    target_font.size = source_font.size
                    target_font.italic = source_font.italic
                    target_font.underline = source_font.underline
                    if source_font.color and source_font.color.rgb:
                        target_font.color.rgb = source_font.color.rgb
            except Exception as e:
                logger.warning(f"[FormFiller] Failed to copy styles: {e}")

    @staticmethod
    def _fill_run_with_value(run, value, pattern=r'_+'):
        """
        Заполнение run'а значением, заменяя подчёркивания.
        Сохраняет форматирование run'а.
        """
        run.text = re.sub(pattern, f' {value}', run.text, count=1)

    @staticmethod
    def _find_underscore_run(paragraph, start_run_index=0):
        """
        Найти первый run с подчёркиваниями начиная с указанного индекса.
        Возвращает индекс run'а или -1.
        """
        pattern = r'_+'
        for i in range(start_run_index, len(paragraph.runs)):
            if re.findall(pattern, paragraph.runs[i].text):
                return i
        return -1

    def fill_and_save(self, row):
        """
        Заполнение форм данными из сформированного в prepare_fill_forms словаря.

        Работает на уровне run'ов и XML для максимального сохранения форматирования.
        Поддерживает field codes (MERGEFIELD, FORMTEXT) и расширенные плейсхолдеры.
        """
        doc = Document(BytesIO(bytes(row.bytes)))
        data = json.loads(row.dictionary)
        pattern = r'_+'
        filled_count = 0
        failed_count = 0

        # Предварительная проверка: есть ли field codes в документе
        has_field_codes = any(
            self._detect_field_codes(p) for p in doc.paragraphs
        )
        if has_field_codes:
            logger.info("[FormFiller] Document contains field codes — will attempt field code filling")

        for key, (ptype, value) in data.items():
            try:
                filled = False

                # Приоритет 0: попробовать заполнить через field codes если они есть
                if has_field_codes and not filled:
                    for paragraph in doc.paragraphs:
                        if self._detect_field_codes(paragraph):
                            if self._clear_text(key) in self._clear_text(paragraph.text):
                                filled = self._fill_field_code(paragraph, str(value))
                                if filled:
                                    logger.info(f"[FormFiller] Filled '{key[:30]}' via field code")
                                    break

                # Стандартное заполнение по типу
                if not filled:
                    if ptype == "Текст до":
                        filled = self._fill_text_before(doc, key, value, pattern)
                    elif ptype == "Текст после":
                        filled = self._fill_text_after(doc, key, value, pattern)
                    elif ptype == "Таблица":
                        filled = self._fill_table(doc, key, value)
                    elif ptype == "Гибрид":
                        filled = self._fill_hybrid(doc, key, value, pattern)

                if filled:
                    filled_count += 1
                else:
                    failed_count += 1
            except Exception as e:
                logger.warning(f"[FormFiller] Failed to fill '{key[:30]}' (type={ptype}): {e}")
                failed_count += 1
                continue

        logger.info(f"[FormFiller] Form filled: {filled_count} ok, {failed_count} failed, total={len(data)}")

        self.replace_last_table_with_podpis(doc)
        buffer = BytesIO()
        doc.save(buffer)
        byte_array = buffer.getvalue()
        buffer.close()
        return list(byte_array)

    def _fill_text_before(self, doc, key, value, pattern):
        """
        Тип 'Текст до': ключевое слово перед пропуском ____ в том же параграфе.
        Работает на уровне run'ов.
        """
        for paragraph in doc.paragraphs:
            if self._clear_text(key) in self._clear_text(paragraph.text):
                for run_number, run in enumerate(paragraph.runs):
                    if self._clear_text(key) in self._clear_text(run.text):
                        key_index = self._clear_text(run.text).find(self._clear_text(key))
                        run_text_full = run.text
                        run_text_slice = run_text_full[key_index + len(key):]

                        if re.findall(pattern, run_text_slice):
                            run_text_slice = re.sub(pattern, f' {value}', run_text_slice, count=1)
                            run.text = run_text_full[:key_index + len(key)] + run_text_slice
                            return True
                        else:
                            # Ищем подчёркивания в следующих run'ах
                            iter_index = 1
                            try:
                                while not re.findall(pattern, paragraph.runs[run_number + iter_index].text):
                                    iter_index += 1
                                target_run = paragraph.runs[run_number + iter_index]
                                target_run.text = re.sub(pattern, f' {value}', target_run.text, count=1)
                                return True
                            except IndexError:
                                continue
                    elif self._clear_text(run.text) in self._clear_text(key) and run.text.strip():
                        try:
                            if re.findall(self._clear_text(paragraph.runs[run_number + 1].text), self._clear_text(key)):
                                continue
                        except IndexError:
                            pass
                        iter_index = 1
                        try:
                            while not re.findall(pattern, paragraph.runs[run_number + iter_index].text):
                                iter_index += 1
                            target_run = paragraph.runs[run_number + iter_index]
                            target_run.text = re.sub(pattern, f' {value}', target_run.text, count=1)
                            return True
                        except IndexError:
                            continue
        return False

    def _fill_text_after(self, doc, key, value, pattern):
        """
        Тип 'Текст после': ключевое слово в скобках после пропуска ____.
        Ищем в предыдущем параграфе.
        """
        for paragraph_number, paragraph in enumerate(doc.paragraphs):
            if self._clear_text(key) in self._clear_text(paragraph.text):
                if paragraph_number > 0:
                    to_replace_paragraph = doc.paragraphs[paragraph_number - 1]
                    for run in to_replace_paragraph.runs[::-1]:
                        run_text = run.text
                        if re.findall(pattern, run_text):
                            run.text = re.sub(pattern, value, run_text, count=1)
                            return True
        return False

    def _fill_table(self, doc, key, value):
        """
        Тип 'Таблица': ключ в ячейке таблицы, значение в соседней (правой) ячейке.
        XML-level копирование rPr для максимального сохранения форматирования.
        """
        for table in doc.tables:
            for row in table.rows:
                for num, cell in enumerate(row.cells):
                    if self._clear_text(key) in self._clear_text(cell.text):
                        try:
                            filling_cell = row.cells[num + 1]
                            if filling_cell.text.strip():
                                # Ячейка не пустая — добавляем через run
                                if filling_cell.paragraphs[-1].runs:
                                    filling_cell.paragraphs[-1].runs[-1].text += "\n" + value
                                else:
                                    filling_cell.paragraphs[-1].add_run("\n" + value)
                            else:
                                # Ячейка пустая — XML-level заполнение
                                self._fill_table_cell_xml(filling_cell, value, source_cell=cell)
                            self.copy_styles(filling_cell.paragraphs[-1], cell.paragraphs[-1])
                            return True
                        except IndexError:
                            # Нет соседней ячейки — добавляем в текущую
                            cell.add_paragraph()
                            new_run = cell.paragraphs[-1].add_run(str(value))
                            if cell.paragraphs[0].runs:
                                source_rpr = cell.paragraphs[0].runs[0]._element.find(qn('w:rPr'))
                                if source_rpr is not None:
                                    new_run._element.insert(0, deepcopy(source_rpr))
                            return True
        return False

    def _fill_hybrid(self, doc, key, value, pattern):
        """
        Тип 'Гибрид': ключ и ____ в одной ячейке таблицы.
        Работает на уровне run'ов и XML для сохранения форматирования.
        НЕ использует cell.text = (уничтожает форматирование).
        """
        # Используем расширенный паттерн плейсхолдеров
        extended_pattern = self.PLACEHOLDER_PATTERN + '|' + pattern if pattern != r'_+' else self.PLACEHOLDER_PATTERN

        for table in doc.tables:
            for row in table.rows:
                for num, cell in enumerate(row.cells):
                    if self._clear_text(key) in self._clear_text(cell.text) and re.findall(extended_pattern, cell.text):
                        # Подчёркивания в той же ячейке что и ключ
                        # Приоритет 1: XML-level через lxml
                        for para in cell.paragraphs:
                            if self._fill_run_xml(para, extended_pattern, f" {value}"):
                                return True
                        # Приоритет 2: через run.text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if re.findall(extended_pattern, run.text):
                                    run.text = re.sub(extended_pattern, f" {value}", run.text, count=1)
                                    return True
                        # НЕ используем cell.text fallback — лучше warning чем потерять форматирование
                        logger.warning(f"[FormFiller] Hybrid: found key '{key[:30]}' but couldn't fill via runs")
                        return False
                    elif self._clear_text(key) in self._clear_text(cell.text) and not re.findall(extended_pattern, cell.text):
                        try:
                            next_cell = row.cells[num + 1]
                            # Приоритет 1: XML-level
                            for para in next_cell.paragraphs:
                                if self._fill_run_xml(para, extended_pattern, f" {value}"):
                                    return True
                            # Приоритет 2: через run.text
                            for para in next_cell.paragraphs:
                                for run in para.runs:
                                    if re.findall(extended_pattern, run.text):
                                        run.text = re.sub(extended_pattern, f" {value}", run.text, count=1)
                                        return True
                            # НЕ используем cell.text fallback
                            logger.warning(f"[FormFiller] Hybrid: found key '{key[:30]}' in adjacent cell but couldn't fill via runs")
                            return False
                        except IndexError:
                            continue
        return False

    def replace_last_table_with_podpis(self, doc):
        tables_with_podpis = []
        paragraphs_with_podpis = []

        # Проходим по всем таблицам в документе
        for i, table in enumerate(doc.tables):
            found = False
            for row in table.rows:
                for cell in row.cells:
                    if 'подпись' in cell.text.lower():
                        found = True
                        break
                if found:
                    break
            if found:
                tables_with_podpis.append(table)

        if not tables_with_podpis:
            to_remove = []
            pattern = r"подпись"
            for num, paragraph in enumerate(doc.paragraphs):
                matches = re.findall(pattern, paragraph.text.lower())
                if matches:
                    paragraphs_with_podpis.append((num, paragraph._element))
            if paragraphs_with_podpis:
                for num, paragraphs_el in paragraphs_with_podpis:
                    try:
                        if "___" in doc.paragraphs[num - 1].text and num - 1 not in [number_par for number_par, _ in to_remove]:
                            to_remove.append((num, doc.paragraphs[num]._element))
                            to_remove.append((num - 1, doc.paragraphs[num - 1]._element))
                        elif "___" in doc.paragraphs[num - 1].text:
                            to_remove.append((num, doc.paragraphs[num]._element))
                            to_remove.append((num + 1, doc.paragraphs[num + 1]._element))
                    except IndexError:
                        pass

                target_paragraph = paragraphs_with_podpis[-1][1]
                new_table = self.create_signature_table(doc)
                par_element = target_paragraph
                par_element.addnext(new_table._tbl)
                for num, paragraph_el in to_remove:
                    paragraph_el.getparent().remove(paragraph_el)
                return None
            else:
                new_table = self.create_signature_table(doc)
                last_el = [cont for cont in doc.iter_inner_content()][-1]
                if last_el._element.tag.endswith('tbl'):
                    tbl_element = last_el._tbl
                    tbl_element.addnext(new_table._tbl)
                elif last_el._element.tag.endswith('p'):
                    par_element = last_el._element
                    par_element.addnext(new_table._tbl)
                return None

        # Выбираем последнюю таблицу
        target_table = tables_with_podpis[-1]

        # Создаём новую таблицу
        new_table = self.create_signature_table(doc)

        # Заменяем старую таблицу на новую
        tbl_element = target_table._tbl
        tbl_element.addnext(new_table._tbl)
        tbl_element.getparent().remove(tbl_element)
        return None

    def create_signature_table(self, doc):
        table = doc.add_table(rows=1, cols=2)
        table.columns[0].width = Pt(300)
        table.columns[1].width = Pt(150)

        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]

        # Левый столбец
        left_cell.paragraphs[0].text = "Начальник Центра торгов"
        left_cell.add_paragraph().text = "Департамента по работе"
        left_cell.add_paragraph().text = "с государственным сектором"

        for par in left_cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.space_before = Pt(0)

        # Правый столбец
        right_cell.add_paragraph()  # index 0
        right_cell.add_paragraph().text = "Баскаков Д.С."  # index 2

        for par in right_cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.space_before = Pt(0)

        # Применяем стиль шрифта
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        return table
